#!/usr/bin/env python3
"""
generate_epsma_docx.py
Generates professional Word (.docx) structured reports adhering to:
E-PSMA: The EANM standardized reporting guidelines v1.0 for PSMA PET (Eur J Nucl Med Mol Imaging 2021 48:1626–1638).

Supports both English (EN) and German (DE) versions.
"""

import sys
import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets background shading of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets inner margins (padding) of a table cell in dxa (1/20 pt)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    """Sets clean subtle borders on a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_epsma_docx(data, output_path, lang="EN"):
    is_de = lang.upper() == "DE"
    doc = Document()

    # Configure Margins: 0.75 in (1.9 cm)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Color Palette
    PRIMARY = RGBColor(27, 54, 93)     # #1B365D Deep Navy
    SECONDARY = RGBColor(70, 130, 180) # Steel Blue
    DARK_TEXT = RGBColor(33, 37, 41)   # Near black
    MUTED = RGBColor(108, 117, 125)    # Gray
    HEADER_BG = "1B365D"               # Navy table header
    ALT_ROW_BG = "F8F9FA"              # Soft light gray zebra

    # Title & Header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_title = title_p.add_run(
        "E-PSMA Strukturierter Befundbericht" if is_de else "E-PSMA Standardized Structured Report"
    )
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    run_sub = sub_p.add_run(
        "Gemäß EANM Standardized Reporting Guidelines v1.0 für PSMA-PET/CT (Eur J Nucl Med Mol Imaging 2021 48:1626–1638)"
        if is_de else
        "In accordance with EANM Standardized Reporting Guidelines v1.0 for PSMA PET/CT (Eur J Nucl Med Mol Imaging 2021 48:1626–1638)"
    )
    run_sub.font.size = Pt(9.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = MUTED

    # Header Meta Box Table
    patient_id = data.get("patient_id", "Unknown")
    exam_date = data.get("study_date", data.get("tp_label", "Current"))
    modality = data.get("modality", "PET/CT")
    mitnm = data.get("final_mitnm", "miT0 miN0 miM0")

    meta_tbl = doc.add_table(rows=2, cols=4)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta_tbl, color="B0C4DE", sz="6")

    meta_headers = ["Patient ID", "Untersuchungsdatum" if is_de else "Exam Date", "Modalität" if is_de else "Modality", "Gesamt miTNM" if is_de else "Overall miTNM"]
    meta_vals = [patient_id, exam_date, modality, mitnm]

    for col_idx, (hdr, val) in enumerate(zip(meta_headers, meta_vals)):
        c0 = meta_tbl.cell(0, col_idx)
        set_cell_background(c0, "E9ECEF")
        set_cell_margins(c0, top=80, bottom=80, left=120, right=120)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(hdr)
        r0.font.size = Pt(8.5)
        r0.font.bold = True
        r0.font.color.rgb = PRIMARY

        c1 = meta_tbl.cell(1, col_idx)
        set_cell_margins(c1, top=80, bottom=80, left=120, right=120)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        r1.font.bold = (col_idx == 3)
        if col_idx == 3:
            r1.font.color.rgb = RGBColor(180, 40, 40)

    # 1. Patient History & Indication
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(4)
    r_h1 = h1.add_run("Klinische Angaben und Anamnese" if is_de else "Patient History")
    r_h1.font.size = Pt(12)
    r_h1.font.bold = True
    r_h1.font.color.rgb = PRIMARY

    hist_txt = data.get("history_text_de" if is_de else "history_text_en", "")
    if not hist_txt:
        hist_txt = data.get("history_text_en", "(No prior clinical indication provided)")
    p_hist = doc.add_paragraph(hist_txt)
    p_hist.paragraph_format.space_after = Pt(10)
    p_hist.style.font.size = Pt(10)

    # 2. Technical information
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    r_h2 = h2.add_run("Untersuchungstechnik" if is_de else "Technical information")
    r_h2.font.size = Pt(12)
    r_h2.font.bold = True
    r_h2.font.color.rgb = PRIMARY

    tech_narr = data.get("tech_narrative_de" if is_de else "tech_narrative_en", "")
    if not tech_narr:
        tech = data.get("tech_params", {})
        inj_act = tech.get("injected_activity", "155 MBq")
        tracer = tech.get("radiotracer", "[68Ga]Ga-PSMA-11")
        up_time = tech.get("uptake_time", "60 minutes")
        ct_prot = tech.get("ct_protocol", "diagnostic")
        if is_de:
            tech_narr = f"Dem Patienten wurden {inj_act} {tracer} intravenös verabreicht. Die Bildgebung erfolgte nach {up_time} mit einer {ct_prot} CT zur Schwächungskorrektur und anatomischen Korrelation."
        else:
            tech_narr = f"The patient was given {inj_act} {tracer} intravenously. Imaging was obtained after {up_time} with {ct_prot} CT for attenuation correction and anatomical correlation."
    
    p_tech = doc.add_paragraph(tech_narr)
    p_tech.paragraph_format.space_after = Pt(10)
    p_tech.style.font.size = Pt(10)

    # 3. Reporting of Findings
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)
    r_h4 = h4.add_run("Befundung" if is_de else "Reporting of Findings")
    r_h4.font.size = Pt(12)
    r_h4.font.bold = True
    r_h4.font.color.rgb = PRIMARY

    preamble_txt = "Dieser Bericht wurde in Übereinstimmung mit den standardisierten Befundungsrichtlinien v1.0 der EANM für PSMA-PET/CT erstellt." if is_de else "This report has been produced in accordance with the European Association of Nuclear Medicine (EANM) standardized reporting guidelines v1.0 for PSMA PET/CT."
    p_pre = doc.add_paragraph(preamble_txt)
    p_pre.paragraph_format.space_after = Pt(10)
    p_pre.style.font.size = Pt(10)

    bio_txt = data.get("biodistribution_text_de" if is_de else "biodistribution_text_en",
                       "Die physiologische Biodistribution des Tracers war regulär." if is_de else
                       "The physiological biodistribution of the radiotracer was regular.")
    p_bio = doc.add_paragraph(bio_txt)
    p_bio.paragraph_format.space_after = Pt(10)
    p_bio.style.font.size = Pt(10)

    regions = [
        ("Prostata" if is_de else "Prostate", data.get("findings_prostate_de" if is_de else "findings_prostate_en", "")),
        ("Lymphknoten" if is_de else "Lymph nodes", data.get("findings_lymph_de" if is_de else "findings_lymph_en", "")),
        ("Skelett" if is_de else "Osseous disease", data.get("findings_bone_de" if is_de else "findings_bone_en", "")),
        ("Viszerale Weichteile" if is_de else "Visceral Disease", data.get("findings_visceral_de" if is_de else "findings_visceral_en", ""))
    ]

    for reg_title, reg_txt in regions:
        if not reg_txt:
            continue
        reg_txt_clean = " ".join([line.strip().lstrip('•>>-').strip() for line in reg_txt.split('\n') if line.strip()])
        p_reg = doc.add_paragraph()
        p_reg.paragraph_format.space_before = Pt(4)
        p_reg.paragraph_format.space_after = Pt(4)
        r_txt = p_reg.add_run(f"- {reg_title}: {reg_txt_clean}")
        r_txt.font.size = Pt(10)
        r_txt.font.color.rgb = DARK_TEXT

    # 4. Conclusion
    h6 = doc.add_paragraph()
    h6.paragraph_format.space_before = Pt(14)
    h6.paragraph_format.space_after = Pt(4)
    r_h6 = h6.add_run("Beurteilung" if is_de else "Conclusion")
    r_h6.font.size = Pt(12)
    r_h6.font.bold = True
    r_h6.font.color.rgb = PRIMARY

    concl_txt = data.get("conclusion_de" if is_de else "conclusion_en", "")
    p_concl = doc.add_paragraph(concl_txt)
    p_concl.paragraph_format.space_after = Pt(10)
    p_concl.style.font.size = Pt(10)

    # Final Staging & Response Callout Box
    tmtv = data.get("tmtv_cc", 0.0)
    recip = data.get("overall_recip", "BASELINE")
    tmtv_delta = data.get("tmtv_delta_pct", 0.0)

    box_tbl = doc.add_table(rows=1, cols=1)
    box_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_box = box_tbl.cell(0, 0)
    set_cell_background(c_box, "F0F4F8")
    set_cell_margins(c_box, top=140, bottom=140, left=200, right=200)
    set_table_borders(box_tbl, color="1B365D", sz="12")

    p_box = c_box.paragraphs[0]
    p_box.paragraph_format.space_after = Pt(0)
    r_b1 = p_box.add_run(f"Zusammenfassendes miTNM-Stadium: {mitnm}\n" if is_de else f"Summary miTNM Classification: {mitnm}\n")
    r_b1.font.bold = True
    r_b1.font.size = Pt(11)
    r_b1.font.color.rgb = PRIMARY

    r_b2 = p_box.add_run(f"• Total Metabolic Tumor Volume (TMTV): {tmtv:.2f} cc\n")
    r_b2.font.size = Pt(9.5)
    if tmtv_delta != 0.0:
        sgn = "+" if tmtv_delta > 0 else ""
        r_b2 = p_box.add_run(f"• Longitudinal TMTV Change: {sgn}{tmtv_delta:.1f}%\n")
        r_b2.font.size = Pt(9.5)

    r_b3 = p_box.add_run(f"• Response Classification (PERCIST/RECIP): {recip}")
    r_b3.font.bold = True
    r_b3.font.size = Pt(9.5)

    # Appendix 
    h_app = doc.add_paragraph()
    h_app.paragraph_format.space_before = Pt(24)
    h_app.paragraph_format.space_after = Pt(10)
    r_app = h_app.add_run("Anhang — Synoptische Tabellen" if is_de else "Appendix — Synoptic Tables")
    r_app.font.size = Pt(14)
    r_app.font.bold = True
    r_app.font.color.rgb = PRIMARY

    # Appendix Table 1
    h2_app = doc.add_paragraph()
    h2_app.paragraph_format.space_before = Pt(10)
    h2_app.paragraph_format.space_after = Pt(4)
    r_h2_app = h2_app.add_run("[Synoptische Tabelle 1 — Untersuchungstechnik]" if is_de else "[Synoptic Table 1 — Technical Parameters]")
    r_h2_app.font.size = Pt(10)
    r_h2_app.font.bold = True
    r_h2_app.font.color.rgb = PRIMARY

    tech = data.get("tech_params", {})
    t1_headers = [
        "Radiopharmakon" if is_de else "Radiotracer",
        "Aktivität" if is_de else "Activity Injected",
        "Uptake-Zeit" if is_de else "Uptake Time",
        "Akquisitionsbereich" if is_de else "Type of Acquisition",
        "CT-Protokoll" if is_de else "CT Protocol",
        "Kontrastmittel" if is_de else "Contrast Oral/IV",
        "Diuretikum" if is_de else "Diuretic"
    ]
    t1_vals = [
        tech.get("radiotracer", "[68Ga]Ga-PSMA-11"),
        tech.get("injected_activity", "155 MBq (2.1 MBq/kg)"),
        tech.get("uptake_time", "60 – 90 min"),
        tech.get("acquisition_type", "Standard (Vertex to mid-thigh)"),
        tech.get("ct_protocol", "Low-Dose / Diagnostic"),
        tech.get("contrast", "No / Nein"),
        tech.get("diuretic", "No / Nein")
    ]

    tbl1 = doc.add_table(rows=2, cols=7)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl1, color="CCCCCC", sz="4")

    for c_idx, (hdr, val) in enumerate(zip(t1_headers, t1_vals)):
        c0 = tbl1.cell(0, c_idx)
        set_cell_background(c0, HEADER_BG)
        set_cell_margins(c0, top=100, bottom=100, left=100, right=100)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(hdr)
        r0.font.size = Pt(8.5)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(255, 255, 255)

        c1 = tbl1.cell(1, c_idx)
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c1, top=80, bottom=80, left=100, right=100)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(val)
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = DARK_TEXT

    # Appendix Table 2
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(14)
    h5.paragraph_format.space_after = Pt(4)
    r_h5 = h5.add_run(
        "[Synoptische Tabelle 2 — Befunde]" if is_de else
        "[Synoptic Table 2 — Findings]"
    )
    r_h5.font.size = Pt(10)
    r_h5.font.bold = True
    r_h5.font.color.rgb = PRIMARY

    t2_headers = [
        "Anatomische Lokalisation" if is_de else "Anatomical Location",
        "miTNM",
        "Größe / Vol." if is_de else "Size / Vol.",
        "Anzahl" if is_de else "Number",
        "PSMA Expr. Q (SUVmax)",
        "PSMA Expr. V",
        "Konfidenz (1-5)" if is_de else "Reader Conf. (1-5)"
    ]

    rows_data = data.get("synoptic_rows", [])
    num_rows = max(len(rows_data), 1) + 1
    tbl2 = doc.add_table(rows=num_rows, cols=7)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl2, color="CCCCCC", sz="4")

    # Header Row
    for c_idx, hdr in enumerate(t2_headers):
        cell = tbl2.cell(0, c_idx)
        set_cell_background(cell, HEADER_BG)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    if not rows_data:
        empty_vals = ["(Keine suspekten Läsionen)" if is_de else "(No suspected lesions identified)", "-", "-", "-", "-", "-", "-"]
        for c_idx, v in enumerate(empty_vals):
            cell = tbl2.cell(1, c_idx)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            run.font.size = Pt(8.5)
    else:
        for r_idx, row in enumerate(rows_data):
            row_num = r_idx + 1
            bg_color = ALT_ROW_BG if (r_idx % 2 == 1) else "FFFFFF"
            r_vals = [
                row.get("location", ""),
                row.get("mitnm", ""),
                row.get("size_str", ""),
                str(row.get("num_lesions", 1)),
                row.get("psma_q", ""),
                row.get("psma_v", ""),
                str(row.get("reader_confidence", 5))
            ]
            for c_idx, val in enumerate(r_vals):
                cell = tbl2.cell(row_num, c_idx)
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                run.font.size = Pt(8.5)
                run.font.color.rgb = DARK_TEXT
                if c_idx == 1:
                    run.font.bold = True

    # Appendix Table 6
    art_rows = data.get("artifact_rows", [])
    if art_rows:
        h_art = doc.add_paragraph()
        h_art.paragraph_format.space_before = Pt(14)
        h_art.paragraph_format.space_after = Pt(4)
        r_hart = h_art.add_run(
            "[Tabelle 6 — Nebenbefunde]" if is_de else
            "[Table 6 — Incidental Findings]"
        )
        r_hart.font.size = Pt(10)
        r_hart.font.bold = True
        r_hart.font.color.rgb = PRIMARY

        t6_headers = [
            "Anatomische Lokalisation" if is_de else "Anatomical Location",
            "Vermutete Ätiologie" if is_de else "Suspected Aetiology",
            "Größe / Vol." if is_de else "Size / Vol.",
            "PSMA Expr. Q (SUVmax)",
            "PSMA Expr. V",
            "Staging-Auswirkung" if is_de else "Staging Impact"
        ]
        tbl6 = doc.add_table(rows=len(art_rows) + 1, cols=6)
        tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl6, color="CCCCCC", sz="4")

        for c_idx, hdr in enumerate(t6_headers):
            cell = tbl6.cell(0, c_idx)
            set_cell_background(cell, "5C6F84")
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(hdr)
            run.font.size = Pt(8.0)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

        for r_idx, row in enumerate(art_rows):
            row_num = r_idx + 1
            bg_color = ALT_ROW_BG if (r_idx % 2 == 1) else "FFFFFF"
            r_vals = [
                row.get("location", ""),
                row.get("comment", "Technical Artifact"),
                row.get("size_str", ""),
                row.get("psma_q", ""),
                row.get("psma_v", ""),
                "Exkludiert (miM0)" if is_de else "Excluded (miM0)"
            ]
            for c_idx, val in enumerate(r_vals):
                cell = tbl6.cell(row_num, c_idx)
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                run.font.size = Pt(8.0)
                run.font.color.rgb = DARK_TEXT
                if c_idx == 5:
                    run.font.bold = True

    # Footnote with reference guidelines
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(18)
    r_ft = p_foot.add_run(
        "Referenzen: (1) Ceci F, et al. E-PSMA: the EANM standardized reporting guidelines v1.0 for PSMA-PET. Eur J Nucl Med Mol Imaging 2021; 48:1626–1638. "
        "(2) Eiber M, et al. Prostate cancer molecular imaging standardized evaluation (PROMISE): proposed miTNM classification. J Nucl Med 2018; 59:469–478."
    )
    r_ft.font.size = Pt(7.5)
    r_ft.font.italic = True
    r_ft.font.color.rgb = MUTED

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"Successfully generated E-PSMA report: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python3 generate_epsma_docx.py <input_json> <output_docx> [EN|DE]")
        sys.exit(1)
    json_path = sys.argv[1]
    out_docx = sys.argv[2]
    lang = sys.argv[3] if len(sys.argv) > 3 else "EN"
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    build_epsma_docx(data, out_docx, lang=lang)
